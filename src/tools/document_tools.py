"""
Universal Document Ingestion & Vector Retrieval Tool for Jarvis.
Supports PDF, Word (.docx), Excel (.xlsx), CSV, TXT, MD, JSON, and code files.
"""

import hashlib
import io
import json
import logging
import sys
import tempfile
import threading
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

# Ensure vendored MinerU is discoverable on sys.path
_MINERU_DIR = Path(__file__).resolve().parent.parent.parent / "MinerU"
if _MINERU_DIR.exists() and str(_MINERU_DIR) not in sys.path:
    sys.path.insert(0, str(_MINERU_DIR))

try:
    from docling.datamodel.base_models import DocumentStream
    from docling.document_converter import DocumentConverter

    DOCLING_AVAILABLE = True
except ImportError:
    DocumentConverter = None  # type: ignore[assignment, misc]
    DocumentStream = None  # type: ignore[assignment, misc]
    DOCLING_AVAILABLE = False

try:
    from langchain_core.tools import BaseTool, create_retriever_tool
except ImportError:
    from langchain.tools.retriever import create_retriever_tool  # type: ignore[no-redef]
    from langchain_core.tools import BaseTool

logger = logging.getLogger(__name__)
logging.getLogger("pypdf").setLevel(logging.ERROR)

_docling_converter_instance: Optional[Any] = None
_docling_lock = threading.Lock()

_MINERU_AVAILABLE: Optional[bool] = None
_mineru_lock = threading.Lock()


def is_docling_available() -> bool:
    """Return whether the Docling document intelligence engine is installed and ready."""
    return DOCLING_AVAILABLE and DocumentConverter is not None


def get_docling_converter() -> Optional[Any]:
    """Retrieve or initialize the thread-safe singleton Docling DocumentConverter."""
    global _docling_converter_instance
    if not is_docling_available():
        return None

    if _docling_converter_instance is None:
        with _docling_lock:
            if _docling_converter_instance is None:
                try:
                    _docling_converter_instance = DocumentConverter()
                    logger.info("Docling DocumentConverter initialized successfully.")
                except Exception as e:
                    logger.warning("Failed to initialize Docling DocumentConverter: %s", e)
                    return None
    return _docling_converter_instance


def convert_document_with_docling(file_name: str, content_bytes: bytes) -> Optional[Tuple[str, Dict[str, Any]]]:
    """
    Attempt to convert a document using Docling into structured Markdown and metadata.
    Returns None if Docling is unavailable or conversion encounters an error.
    """
    converter = get_docling_converter()
    if converter is None or DocumentStream is None:
        return None

    try:
        stream = DocumentStream(name=file_name, stream=io.BytesIO(content_bytes))
        result = converter.convert(stream)
        if result and result.document:
            markdown = result.document.export_to_markdown()
            meta: Dict[str, Any] = {
                "filename": file_name,
                "type": Path(file_name).suffix.lower(),
                "size": len(content_bytes),
                "engine": "docling",
                "tables": len(result.document.tables) if hasattr(result.document, "tables") else 0,
            }
            if hasattr(result.document, "pages"):
                meta["pages"] = len(result.document.pages)
            return markdown, meta
    except Exception as e:
        logger.debug("Docling conversion skipped for %s, falling back to native extractor: %s", file_name, e)
    return None


def is_mineru_available() -> bool:
    """Return whether the MinerU document intelligence engine is installed and ready."""
    global _MINERU_AVAILABLE
    if _MINERU_AVAILABLE is None:
        try:
            import mineru  # noqa: F401

            _MINERU_AVAILABLE = True
        except ImportError:
            _MINERU_AVAILABLE = False
    return _MINERU_AVAILABLE


def _count_mineru_features(pdf_info: List[Dict[str, Any]]) -> Tuple[int, int]:
    """Count LaTeX formulas and tables in MinerU pdf_info data structure."""
    formulas = 0
    tables = 0
    for page in pdf_info:
        for block in page.get("para_blocks", []):
            btype = str(block.get("type", "")).lower()
            if "equation" in btype or "formula" in btype:
                formulas += 1
            elif "table" in btype:
                tables += 1
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    stype = str(span.get("type", "")).lower()
                    content = str(span.get("content", ""))
                    if "equation" in stype or "formula" in stype:
                        formulas += 1
                    elif "$" in content and not ("equation" in btype or "formula" in btype) and content.count("$") >= 2:
                        formulas += 1
    return formulas, tables


def _run_mineru_pdf_extraction(
    file_name: str,
    content_bytes: bytes,
    parse_method: str = "auto",
    formula_enable: bool = True,
    table_enable: bool = True,
) -> Optional[Tuple[str, Dict[str, Any]]]:
    """Execute MinerU PDF extraction pipeline."""
    try:
        from mineru.backend.pipeline.pipeline_analyze import doc_analyze_streaming
        from mineru.backend.pipeline.pipeline_middle_json_mkcontent import union_make as pipeline_union_make
        from mineru.data.data_reader_writer import FileBasedDataWriter
        from mineru.utils.enum_class import MakeMode

        with tempfile.TemporaryDirectory() as temp_dir:
            image_writer = FileBasedDataWriter(str(Path(temp_dir) / "images"))
            ready_docs: List[Dict[str, Any]] = []

            def on_doc_ready(doc_index: int, model_list: Any, middle_json: Dict[str, Any], ocr_enable: bool) -> None:
                ready_docs.append(middle_json)

            doc_analyze_streaming(
                pdf_bytes_list=[content_bytes],
                image_writer_list=[image_writer],
                lang_list=["en"],
                on_doc_ready=on_doc_ready,
                parse_method=parse_method,
                formula_enable=formula_enable,
                table_enable=table_enable,
            )

            if ready_docs and "pdf_info" in ready_docs[0]:
                middle_json = ready_docs[0]
                pdf_info = middle_json["pdf_info"]
                md_text = pipeline_union_make(pdf_info, MakeMode.MM_MD)
                if md_text and md_text.strip():
                    num_formulas, num_tables = _count_mineru_features(pdf_info)
                    meta: Dict[str, Any] = {
                        "filename": file_name,
                        "type": ".pdf",
                        "size": len(content_bytes),
                        "engine": "mineru",
                        "pages": len(pdf_info),
                        "formulas": num_formulas,
                        "tables": num_tables,
                    }
                    return md_text, meta
    except Exception as e:
        logger.debug("MinerU PDF extraction encountered error for %s: %s", file_name, e)
    return None


def _run_mineru_office_extraction(file_name: str, content_bytes: bytes) -> Optional[Tuple[str, Dict[str, Any]]]:
    """Execute MinerU Office (DOCX, PPTX, XLSX) extraction."""
    suffix = Path(file_name).suffix.lower()
    try:
        from mineru.backend.office.office_middle_json_mkcontent import union_make as office_union_make
        from mineru.utils.enum_class import MakeMode

        middle_json = None
        if suffix == ".docx":
            from mineru.backend.office.docx_analyze import office_docx_analyze

            middle_json, _ = office_docx_analyze(content_bytes)
        elif suffix == ".pptx":
            from mineru.backend.office.pptx_analyze import office_pptx_analyze

            middle_json, _ = office_pptx_analyze(content_bytes)
        elif suffix == ".xlsx":
            from mineru.backend.office.xlsx_analyze import office_xlsx_analyze

            middle_json, _ = office_xlsx_analyze(content_bytes)

        if middle_json and "pdf_info" in middle_json:
            pdf_info = middle_json["pdf_info"]
            md_text = office_union_make(pdf_info, MakeMode.MM_MD)
            if md_text and md_text.strip():
                num_formulas, num_tables = _count_mineru_features(pdf_info)
                meta: Dict[str, Any] = {
                    "filename": file_name,
                    "type": suffix,
                    "size": len(content_bytes),
                    "engine": "mineru",
                    "pages": len(pdf_info),
                    "formulas": num_formulas,
                    "tables": num_tables,
                }
                return md_text, meta
    except Exception as e:
        logger.debug("MinerU office extraction encountered error for %s: %s", file_name, e)
    return None


def convert_document_with_mineru(
    file_name: str,
    content_bytes: bytes,
    parse_method: str = "auto",
    formula_enable: bool = True,
    table_enable: bool = True,
) -> Optional[Tuple[str, Dict[str, Any]]]:
    """
    Attempt to convert a document using MinerU into structured Markdown and metadata.
    Specializes in academic/scientific PDFs, LaTeX formula recognition ($...$ / $$...$$),
    multi-column reading order, and complex tables.
    Returns (markdown, metadata) or None if MinerU is unavailable or encounters an error.
    """
    if not is_mineru_available():
        return None

    suffix = Path(file_name).suffix.lower()
    if suffix not in [".pdf", ".docx", ".pptx", ".xlsx"]:
        return None

    # 1. Custom runner hook (allows unit test mocking and custom pipeline injection)
    if hasattr(convert_document_with_mineru, "_runner"):
        runner = convert_document_with_mineru._runner
        if callable(runner):
            try:
                return runner(
                    file_name,
                    content_bytes,
                    parse_method=parse_method,
                    formula_enable=formula_enable,
                    table_enable=table_enable,
                )
            except Exception as e:
                logger.debug("MinerU runner encountered error for %s: %s", file_name, e)
                return None

    # 2. Native execution with thread safety
    with _mineru_lock:
        try:
            if suffix == ".pdf":
                return _run_mineru_pdf_extraction(
                    file_name,
                    content_bytes,
                    parse_method=parse_method,
                    formula_enable=formula_enable,
                    table_enable=table_enable,
                )
            elif suffix in [".docx", ".pptx", ".xlsx"]:
                return _run_mineru_office_extraction(file_name, content_bytes)
        except Exception as e:
            logger.debug("MinerU conversion skipped for %s, falling back to next engine: %s", file_name, e)
    return None


def parse_document_with_mineru(file_obj: Any, formula_enable: bool = True, table_enable: bool = True) -> Dict[str, Any]:
    """
    Parse a document using MinerU document intelligence engine with deep formula and table extraction.
    Returns a dictionary with extracted markdown, metadata, and status.
    """
    file_name = file_obj.name
    content_bytes = file_obj.getvalue() if hasattr(file_obj, "getvalue") else file_obj.read()
    if hasattr(file_obj, "seek"):
        file_obj.seek(0)

    res = convert_document_with_mineru(
        file_name=file_name,
        content_bytes=content_bytes,
        formula_enable=formula_enable,
        table_enable=table_enable,
    )
    if res is not None:
        text, meta = res
        return {"status": "success", "text": text, "metadata": meta, "engine": "mineru"}

    text, meta = extract_text_from_file(file_obj)
    return {"status": "fallback", "text": text, "metadata": meta, "engine": meta.get("engine", "native")}


def create_mineru_document_tool() -> BaseTool:
    """Build a LangChain tool for high-precision scientific and academic document parsing via MinerU."""
    from langchain_core.tools import tool

    @tool
    def parse_scientific_document(file_path: str) -> str:
        """Parse a scientific or academic document (PDF, DOCX) extracting LaTeX formulas ($...$ / $$...$$), tables, and structured markdown using MinerU."""
        path = Path(file_path)
        if not path.exists():
            return f"Error: File not found at {file_path}"
        with open(path, "rb") as f:
            content = f.read()
        res = convert_document_with_mineru(path.name, content)
        if res is not None:
            text, meta = res
            return f"=== MinerU Document Analysis: {path.name} (Formulas: {meta.get('formulas', 0)}, Tables: {meta.get('tables', 0)}) ===\n\n{text}"

        class _MockFile:
            name = path.name

            def read(self) -> bytes:
                return content

            def getvalue(self) -> bytes:
                return content

        text, meta = extract_text_from_file(_MockFile())
        return f"=== Document Content: {path.name} ===\n\n{text}"

    return parse_scientific_document


def get_files_hash(uploaded_files: List[Any]) -> str:
    """Generate a combined SHA-256 hash of all uploaded files for caching."""
    hasher = hashlib.sha256()
    for file in uploaded_files:
        content = file.getvalue() if hasattr(file, "getvalue") else file.read()
        hasher.update(file.name.encode("utf-8"))
        hasher.update(content)
        if hasattr(file, "seek"):
            file.seek(0)
    return hasher.hexdigest()


def extract_text_from_file(file: Any, prefer_engine: Optional[str] = None) -> Tuple[str, Dict[str, Any]]:
    """
    Extract readable text and metadata from a single uploaded file.
    Supports PDF, DOCX, PPTX, CSV, XLSX, TXT, MD, HTML, JSON, PY.
    Leverages MinerU (specialized in academic PDFs, LaTeX formulas, complex tables)
    and Docling (universal structured document layouts) with automatic fallback to native parsers.
    """
    file_name = file.name
    suffix = Path(file_name).suffix.lower()
    text = ""
    content_bytes = file.getvalue() if hasattr(file, "getvalue") else file.read()
    metadata: Dict[str, Any] = {"filename": file_name, "type": suffix, "size": len(content_bytes)}

    # 1. Attempt MinerU conversion first when requested or for PDF files by default
    if suffix in [".pdf", ".docx", ".pptx", ".xlsx"] and (
        prefer_engine == "mineru" or (prefer_engine is None and suffix == ".pdf" and is_mineru_available())
    ):
        mineru_result = convert_document_with_mineru(file_name, content_bytes)
        if mineru_result is not None:
            mineru_text, mineru_meta = mineru_result
            if mineru_text.strip():
                return mineru_text, mineru_meta

    # 2. Attempt Docling conversion for rich structured document formats
    if suffix in [".pdf", ".docx", ".pptx", ".html", ".md"] and prefer_engine != "native":
        docling_result = convert_document_with_docling(file_name, content_bytes)
        if docling_result is not None:
            docling_text, docling_meta = docling_result
            if docling_text.strip():
                return docling_text, docling_meta

    # 3. If MinerU was not attempted yet (e.g. non-PDF or prefer_engine was docling and docling failed)
    if suffix in [".pdf", ".docx", ".pptx", ".xlsx"] and prefer_engine not in ["native", "mineru"]:
        mineru_result = convert_document_with_mineru(file_name, content_bytes)
        if mineru_result is not None:
            mineru_text, mineru_meta = mineru_result
            if mineru_text.strip():
                return mineru_text, mineru_meta

    try:
        # Fallback to native extractors
        # 1. PDF fallback
        if suffix == ".pdf":
            pdf_reader = PdfReader(io.BytesIO(content_bytes))
            num_pages = len(pdf_reader.pages)
            metadata["pages"] = num_pages
            page_texts = []
            for idx, page in enumerate(pdf_reader.pages):
                extracted = page.extract_text()
                if extracted:
                    page_texts.append(f"--- [Page {idx + 1} of {file_name}] ---\n{extracted}")
            text = "\n\n".join(page_texts)

        # 2. DOCX (Word)
        elif suffix == ".docx":
            try:
                import docx

                doc = docx.Document(io.BytesIO(content_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                metadata["paragraphs"] = len(paragraphs)
                text = f"=== Word Document: {file_name} ===\n" + "\n\n".join(paragraphs)
            except ImportError:
                text = f"[Warning: python-docx not installed, reading as raw text]\n{content_bytes.decode('utf-8', errors='ignore')}"

        # 3. CSV
        elif suffix == ".csv":
            try:
                df = pd.read_csv(io.BytesIO(content_bytes))
                metadata["rows"] = len(df)
                metadata["columns"] = list(df.columns)
                summary_str = f"=== CSV Dataset: {file_name} ===\nShape: {df.shape[0]} rows x {df.shape[1]} columns\n"
                summary_str += f"Columns: {', '.join(df.columns)}\n\nFirst 25 Rows:\n{df.head(25).to_markdown(index=False)}\n\nStatistical Summary:\n{df.describe(include='all').to_string()}"
                text = summary_str
            except Exception:
                text = content_bytes.decode("utf-8", errors="ignore")

        # 4. Excel (XLSX, XLS)
        elif suffix in [".xlsx", ".xls"]:
            try:
                xls = pd.ExcelFile(io.BytesIO(content_bytes))
                metadata["sheets"] = xls.sheet_names
                sheets_text = [f"=== Excel Workbook: {file_name} ==="]
                for sheet in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet)
                    sheets_text.append(
                        f"--- Sheet: {sheet} ({df.shape[0]} rows x {df.shape[1]} cols) ---\n{df.head(20).to_markdown(index=False)}"
                    )
                text = "\n\n".join(sheets_text)
            except Exception as e:
                text = f"Error reading excel: {str(e)}"

        # 5. JSON
        elif suffix == ".json":
            data = json.loads(content_bytes.decode("utf-8", errors="ignore"))
            text = f"=== JSON Document: {file_name} ===\n" + json.dumps(data, indent=2)

        # 6. TXT, Markdown, Python, Code
        else:
            decoded = content_bytes.decode("utf-8", errors="ignore")
            text = f"=== Document: {file_name} ===\n{decoded}"

    except Exception as e:
        logger.error(f"Error extracting text from {file_name}: {str(e)}", exc_info=True)
        text = f"Error reading {file_name}: {str(e)}"

    return text, metadata


def process_documents_and_build_vector_store(
    uploaded_files: List[Any], api_provider: str = "OpenRouter", chunk_size: int = 1000, chunk_overlap: int = 150
) -> Tuple[Optional[FAISS], List[Dict[str, Any]], str]:
    """
    Process all uploaded document files, extract text, split into chunks,
    and construct a FAISS vector store.
    """
    if not uploaded_files:
        return None, [], ""

    all_texts = []
    file_summaries = []

    for file in uploaded_files:
        doc_text, meta = extract_text_from_file(file)
        if doc_text.strip():
            all_texts.append(doc_text)
            file_summaries.append(meta)

    if not all_texts:
        return None, file_summaries, "No readable text could be extracted from uploaded files."

    combined_text = "\n\n" + "=" * 50 + "\n\n".join(all_texts)

    # Chunking
    text_splitter = RecursiveCharacterTextSplitter(
        separators=["\n\n", "\n", " ", ""], chunk_size=chunk_size, chunk_overlap=chunk_overlap, length_function=len
    )
    chunks = text_splitter.split_text(combined_text)
    logger.info(f"Split documents into {len(chunks)} chunks.")

    # Embeddings
    try:
        if api_provider == "OpenAI":
            embeddings: Any = OpenAIEmbeddings()
        else:
            embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

        vector_store = FAISS.from_texts(chunks, embeddings)
        return (
            vector_store,
            file_summaries,
            f"Successfully processed {len(uploaded_files)} file(s) into {len(chunks)} searchable vectors.",
        )
    except Exception as e:
        logger.error(f"Vector store creation error: {str(e)}", exc_info=True)
        return None, file_summaries, f"Embedding creation failed: {str(e)}"


def create_document_retriever_tool(vector_store: FAISS, top_k: int = 4) -> BaseTool:
    """Build a LangChain retriever tool from the FAISS vector store."""
    retriever = vector_store.as_retriever(search_kwargs={"k": top_k})
    return create_retriever_tool(
        retriever,
        "document_search",
        "Searches and retrieves the most relevant excerpts, tables, and sections from all uploaded files (PDF, Word, Excel, CSV, text).",
    )


def extract_entities_from_document(
    file_obj: Any,
    prompt_description: str = "Extract key entities, facts, and attributes",
    model_id: str = "gpt-4o-mini",
    **kwargs: Any,
) -> Dict[str, Any]:
    """
    Extract structured, grounded entities from an uploaded document.
    Leverages Docling for deep document layout parsing, then runs LangExtract
    to map entities to exact character spans in the extracted text.
    """
    from .extraction_tools import extract_grounded_entities

    text, meta = extract_text_from_file(file_obj)
    extraction_res = extract_grounded_entities(
        text=text,
        prompt_description=prompt_description,
        model_id=model_id,
        **kwargs,
    )
    extraction_res["document_metadata"] = meta
    return extraction_res
