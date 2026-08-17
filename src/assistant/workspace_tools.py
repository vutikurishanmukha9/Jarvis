"""
Workspace Management and Artifact Generation Tools for Auto-JARVIS.
Enables autonomous file operations (read, write, edit, list) and artifact creation (.xlsx, .docx, .md, .csv).
"""

import json
import logging
import os
from pathlib import Path
from typing import Dict, Any, List, Optional
import pandas as pd
from langchain_core.tools import tool, BaseTool

from ..config import WORKSPACE_DIR
from .profile_manager import ProfileManager

logger = logging.getLogger(__name__)

def _resolve_workspace_path(path_str: str) -> Path:
    """Resolve a relative path, rejecting every attempt to leave ``WORKSPACE_DIR``."""
    if not isinstance(path_str, str) or not path_str.strip():
        raise ValueError("A non-empty workspace-relative path is required.")

    candidate = Path(path_str.strip())
    if candidate.is_absolute():
        raise ValueError("Absolute paths are not allowed in the workspace.")

    workspace_resolved = WORKSPACE_DIR.resolve()
    target = (workspace_resolved / candidate).resolve()
    try:
        target.relative_to(workspace_resolved)
    except ValueError as exc:
        raise ValueError("Path must remain inside the workspace.") from exc
    return target

@tool
def write_workspace_file(filename: str, content: str) -> str:
    """
    Creates or overwrites a file in the workspace sandbox.
    Supports .md, .txt, .py, .csv, .json, .html, etc.
    Use this tool whenever you need to produce a report, write a script, generate a dataset, or save notes on the user's behalf.
    """
    try:
        target = _resolve_workspace_path(filename)
        target.parent.mkdir(parents=True, exist_ok=True)
        with open(target, "w", encoding="utf-8") as f:
            f.write(content)
        size_bytes = len(content.encode("utf-8"))
        logger.info(f"Wrote workspace file: {target.name} ({size_bytes} bytes)")
        return f"Successfully created workspace file: '{target.name}' ({size_bytes} bytes) at {target}"
    except Exception as e:
        logger.error(f"Failed to write file {filename}: {str(e)}")
        return f"Error writing file '{filename}': {str(e)}"

@tool
def read_workspace_file(filename: str) -> str:
    """
    Reads the content of an existing file in the workspace sandbox.
    Use this to inspect files previously created or uploaded into the workspace.
    """
    try:
        target = _resolve_workspace_path(filename)
        if not target.exists():
            return f"File '{filename}' does not exist in workspace."
        with open(target, "r", encoding="utf-8") as f:
            content = f.read()
        return f"=== Content of '{target.name}' ({len(content)} chars) ===\n{content}"
    except Exception as e:
        return f"Error reading file '{filename}': {str(e)}"

@tool
def list_workspace_files(subdirectory: str = "") -> str:
    """
    Lists all files and directories currently inside the workspace.
    Use this tool to check existing deliverables, dataset files, or documents created for the user.
    """
    try:
        target_dir = _resolve_workspace_path(subdirectory) if subdirectory else WORKSPACE_DIR
        if not target_dir.exists():
            return "Workspace directory is currently empty."
        
        files = list(target_dir.rglob("*"))
        if not files:
            return "Workspace is empty. No files created yet."
        
        file_list = []
        for p in sorted(files):
            rel_path = p.relative_to(WORKSPACE_DIR)
            if p.is_file():
                size_kb = round(p.stat().st_size / 1024, 2)
                file_list.append(f"- [FILE] {rel_path} ({size_kb} KB)")
            elif p.is_dir():
                file_list.append(f"- [DIR]  {rel_path}/")
        
        return "=== Workspace Files ===\n" + "\n".join(file_list)
    except Exception as e:
        return f"Error listing workspace: {str(e)}"

@tool
def generate_excel_spreadsheet(filename: str, json_table_data: str, sheet_name: str = "Data") -> str:
    """
    Generates a structured Microsoft Excel (.xlsx) spreadsheet in the workspace.
    `json_table_data` must be a JSON array of objects representing rows, e.g. '[{"Name": "A", "Revenue": 100}, {"Name": "B", "Revenue": 200}]'.
    Use this tool when the user asks to build financial models, data comparison tables, or structured spreadsheets.
    """
    try:
        if not filename.endswith(".xlsx"):
            filename += ".xlsx"
        target = _resolve_workspace_path(filename)
        target.parent.mkdir(parents=True, exist_ok=True)
        
        records = json.loads(json_table_data)
        df = pd.DataFrame(records)
        df.to_excel(target, index=False, sheet_name=sheet_name)
        
        return f"Successfully generated Excel spreadsheet: '{target.name}' with {len(df)} rows and {len(df.columns)} columns."
    except Exception as e:
        logger.error(f"Excel generation failed: {str(e)}")
        return f"Error generating Excel file: {str(e)}"

@tool
def generate_word_document(filename: str, title: str, markdown_content: str) -> str:
    """
    Generates a formatted Microsoft Word (.docx) document in the workspace from Markdown text.
    Use this tool when the user asks for a formal report, whitepaper, project brief, or proposal.
    """
    try:
        if not filename.endswith(".docx"):
            filename += ".docx"
        target = _resolve_workspace_path(filename)
        target.parent.mkdir(parents=True, exist_ok=True)

        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            doc = Document()
            
            # Title
            title_heading = doc.add_heading(title, level=0)
            
            # Parse Markdown paragraphs and headings
            lines = markdown_content.split("\n")
            for line in lines:
                s_line = line.strip()
                if not s_line:
                    continue
                if s_line.startswith("### "):
                    doc.add_heading(s_line[4:], level=2)
                elif s_line.startswith("## "):
                    doc.add_heading(s_line[3:], level=1)
                elif s_line.startswith("# "):
                    doc.add_heading(s_line[2:], level=1)
                elif s_line.startswith("- ") or s_line.startswith("* "):
                    doc.add_paragraph(s_line[2:], style='List Bullet')
                elif s_line[0:2].isdigit() and s_line[2:4] in [". ", ") "]:
                    doc.add_paragraph(s_line[3:], style='List Number')
                else:
                    doc.add_paragraph(s_line)
                    
            doc.save(target)
            return f"Successfully generated Word Document: '{target.name}' at {target}"
        except ImportError:
            # Fallback to saving markdown if docx library is not available
            fallback_target = target.with_suffix(".md")
            with open(fallback_target, "w", encoding="utf-8") as f:
                f.write(f"# {title}\n\n{markdown_content}")
            return f"Saved document as Markdown '{fallback_target.name}' (python-docx not installed)."
    except Exception as e:
        logger.error(f"Word document generation failed: {str(e)}")
        return f"Error creating Word document: {str(e)}"

@tool
def save_personal_memory(fact: str, category: str = "preference") -> str:
    """
    Saves a persistent long-term memory fact about the user, their projects, or preferences.
    Use this tool whenever the user tells you to remember something for future tasks.
    """
    success = ProfileManager.add_memory(fact, category)
    if success:
        return f"Logged into persistent memory: '{fact}' [Category: {category}]"
    return "Failed to record memory entry."

def get_workspace_tools() -> List[BaseTool]:
    """Retrieve the full suite of workspace and assistant automation tools."""
    return [
        write_workspace_file,
        read_workspace_file,
        list_workspace_files,
        generate_excel_spreadsheet,
        generate_word_document,
        save_personal_memory
    ]
